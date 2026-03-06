import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import io
import pandas as pd
import zipfile
import re
import base64 
from pyproj import Transformer
from datetime import datetime
import locale

# --- IMPORTACIÓN NUEVA PARA PDF ---
try:
    import fitz  # PyMuPDF
except ImportError:
    st.error("⚠️ Falta instalar la librería 'pymupdf'. Agregala a requirements.txt")

# --- IMPORTACIONES PARA MAPA ---
try:
    import folium
    from streamlit_folium import st_folium
except ImportError:
    pass

# --- CONFIGURACIÓN GLOBAL ---
st.set_page_config(page_title="Arqueología - Suite Word", layout="wide")

# ==========================================
# 1. FUNCIONES AUXILIARES
# ==========================================

def obtener_imagenes_con_id(elemento_xml, doc_relacionado):
    """Extrae imágenes incrustadas en una celda/párrafo de Word."""
    resultados = [] 
    blips = elemento_xml.xpath('.//a:blip')
    for blip in blips:
        try:
            embed_code = blip.get(qn('r:embed'))
            if embed_code:
                part = doc_relacionado.part.related_parts[embed_code]
                if 'image' in part.content_type:
                    resultados.append((embed_code, part.blob))
        except:
            continue
    return resultados

def obtener_texto_celda_abajo(tabla, fila_idx, col_idx):
    try:
        if fila_idx + 1 < len(tabla.rows):
            fila_siguiente = tabla.rows[fila_idx + 1]
            if col_idx < len(fila_siguiente.cells):
                return fila_siguiente.cells[col_idx].text.strip()
    except:
        pass
    return ""

def limpiar_coordenada(texto):
    texto_limpio = texto.replace(".", "").replace(" ", "").strip()
    texto_limpio = texto_limpio.replace(",", ".")
    try:
        return float(texto_limpio)
    except:
        return None

# ==========================================
# 2. LÓGICA: GENERADOR WORD (MAP - DESDE WORD)
# ==========================================

def procesar_archivo_v12(archivo_bytes, nombre_archivo):
    try:
        doc = Document(io.BytesIO(archivo_bytes))
    except Exception as e:
        st.error(f"Error leyendo {nombre_archivo}: {e}")
        return []

    fichas_extraidas = []
    fecha_persistente = "Sin Fecha"

    for tabla in doc.tables:
        datos_ficha = {
            "fecha_propia": None, "actividad": "", "hallazgos": "", "items_foto": [] 
        }
        rids_procesados = set()
        celdas_procesadas = set()
        en_seccion_fotos = False
        
        for r_idx, fila in enumerate(tabla.rows):
            texto_fila = " ".join([c.text.strip() for c in fila.cells]).strip()
            
            if "Fecha" in texto_fila:
                for celda in fila.cells:
                    t = celda.text.strip()
                    if "Fecha" not in t and len(t) > 5:
                        datos_ficha["fecha_propia"] = t
                        fecha_persistente = t
                        break
            
            if "Descripción de la actividad" in texto_fila:
                mejor_texto = ""
                celdas_fila_vistas = set()
                for celda in fila.cells:
                    if celda in celdas_fila_vistas: continue
                    celdas_fila_vistas.add(celda)
                    t = celda.text.strip()
                    if "Descripción" in t or "Actividad" in t: continue
                    if len(t) > len(mejor_texto):
                        mejor_texto = t
                if mejor_texto:
                    datos_ficha["actividad"] = mejor_texto

            if "Ausencia" in texto_fila and any(c.text.strip().upper() == "X" for c in fila.cells):
                datos_ficha["hallazgos"] = "Ausencia de hallazgos arqueológicos no previstos."
            if "Presencia" in texto_fila and any(c.text.strip().upper() == "X" for c in fila.cells):
                datos_ficha["hallazgos"] = "PRESENCIA de hallazgos arqueológicos."

            if "Registro fotográfico" in texto_fila:
                en_seccion_fotos = True
                continue 

            if en_seccion_fotos:
                for c_idx, celda in enumerate(fila.cells):
                    if celda in celdas_procesadas: continue
                    celdas_procesadas.add(celda)

                    lista_imgs_ids = obtener_imagenes_con_id(celda._element, doc)
                    if lista_imgs_ids:
                        texto_leyenda = celda.text.strip()
                        if not texto_leyenda:
                            texto_leyenda = obtener_texto_celda_abajo(tabla, r_idx, c_idx)
                        
                        for rId, blob in lista_imgs_ids:
                            if rId in rids_procesados: continue
                            rids_procesados.add(rId)
                            datos_ficha["items_foto"].append({
                                "blob": blob, "leyenda": texto_leyenda
                            })
                celdas_procesadas.clear() 

        fecha_final = datos_ficha["fecha_propia"] if datos_ficha["fecha_propia"] else fecha_persistente
        
        if datos_ficha["actividad"] or datos_ficha["items_foto"]:
            texto_central = datos_ficha["actividad"]
            if datos_ficha["hallazgos"]:
                texto_central += f"\n\n[Hallazgos: {datos_ficha['hallazgos']}]"
            
            fichas_extraidas.append({
                "fecha": fecha_final, "texto_central": texto_central, "fotos": datos_ficha["items_foto"]
            })

    return fichas_extraidas

def generar_word_con_formato(datos):
    doc = Document()
    titulo = doc.add_heading('Tabla Resumen Monitoreo Arqueológico', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'
    tabla.autofit = False
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER 

    headers = tabla.rows[0].cells
    titulos = ["Fecha", "Actividades realizadas durante el MAP", "Imagen de la actividad"]
    
    for i, texto in enumerate(titulos):
        parrafo = headers[i].paragraphs[0]
        run = parrafo.add_run(texto)
        run.font.name = 'Franklin Gothic Book'
        run.font.size = Pt(9)
        run.bold = True
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for c in tabla.columns[0].cells: c.width = Cm(2.5) 
    for c in tabla.columns[1].cells: c.width = Cm(7.5) 
    for c in tabla.columns[2].cells: c.width = Cm(8.5) 

    for item in datos:
        row = tabla.add_row().cells
        p_fecha = row[0].paragraphs[0]
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_fecha = p_fecha.add_run(str(item["fecha"]))
        r_fecha.font.name = 'Franklin Gothic Book'
        r_fecha.font.size = Pt(9)

        p_act = row[1].paragraphs[0]
        p_act.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_act = p_act.add_run(str(item["texto_central"]))
        r_act.font.name = 'Franklin Gothic Book'
        r_act.font.size = Pt(9)
        
        celda_img = row[2]
        p_img = celda_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if not item["fotos"]:
            r_sin = p_img.add_run("[Sin fotos]")
            r_sin.font.name = 'Franklin Gothic Book'
            r_sin.font.size = Pt(9)
        else:
            for i, foto_obj in enumerate(item["fotos"]):
                try:
                    run = p_img.add_run()
                    run.add_picture(io.BytesIO(foto_obj["blob"]), width=Cm(8), height=Cm(6))
                    if foto_obj["leyenda"]:
                        r_leyenda = p_img.add_run(f"\n{foto_obj['leyenda']}")
                        r_leyenda.font.name = 'Franklin Gothic Book'
                        r_leyenda.font.size = Pt(9)
                        r_leyenda.italic = True
                    if i < len(item["fotos"]) - 1:
                        p_img.add_run("\n\n")
                except:
                    continue
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 2.1 LÓGICA NUEVA: GENERADOR WORD MAP (DESDE PDF) - V7 MULTIPÁGINA
# ==========================================

def procesar_pdf_a_word_map(pdf_bytes, nombre_archivo):
    """
    Extrae Fecha, Actividad y Fotos de reportes en PDF usando PyMuPDF (fitz).
    LÓGICA V7 (Estado Persistente):
    - Activa captura cuando encuentra "V. DESCRIPCIONES".
    - Mantiene captura activa a través de saltos de página.
    - Desactiva captura cuando encuentra "VI. CARACTERÍSTICAS".
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        st.error(f"Error abriendo PDF {nombre_archivo}: {e}")
        return []

    fichas = []
    ficha_actual = {
        "fecha": None,
        "texto_central": "",
        "fotos": []
    }
    
    # ESTADO PERSISTENTE ENTRE PÁGINAS
    capturando_descripcion = False
    
    # Textos a limpiar del resultado final
    blacklist_clean = [
        "V. DESCRIPCIONES", "Descripción de la Actividad", 
        "Huso", "18 G", "19 H", "Datum", "WGS84",
        "Coordenadas", "Vértice", "Este", "Norte", "Altitud"
    ]

    for pagina_idx, pagina in enumerate(doc):
        # 1. ORDENAR BLOQUES VISUALMENTE
        bloques = pagina.get_text("blocks")
        bloques.sort(key=lambda b: (b[1], b[0])) 
        
        texto_plano_pagina = pagina.get_text("text")

        # DETECTAR NUEVA FICHA (Reset general)
        if "I. IDENTIFICACIÓN" in texto_plano_pagina or "Ficha de Monitoreo Arqueológico" in texto_plano_pagina:
            if ficha_actual["fecha"] or ficha_actual["texto_central"] or ficha_actual["fotos"]:
                fichas.append(ficha_actual)
            ficha_actual = { "fecha": None, "texto_central": "", "fotos": [] }
            capturando_descripcion = False # Seguridad: Resetear estado al iniciar nueva ficha

        # 2. EXTRAER FECHA
        if not ficha_actual["fecha"]:
            match_fecha = re.search(r"(\d{2}/\d{2}/\d{4})", texto_plano_pagina)
            if match_fecha:
                ficha_actual["fecha"] = match_fecha.group(1)

        # 3. EXTRAER ACTIVIDAD (Lógica de Estado Persistente)
        for b in bloques:
            txt = b[4].strip()
            
            # A. DETECTAR INICIO
            # Si encontramos el título de inicio, activamos bandera
            if "V. DESCRIPCIONES" in txt or "Descripción de la Actividad" in txt:
                capturando_descripcion = True
                continue # No guardamos el título en sí

            # B. DETECTAR FIN
            # Si encontramos el título de fin, apagamos bandera
            if "VI. CARACTERÍSTICAS" in txt or "CARACTERÍSTICAS DE LA CAPA" in txt:
                capturando_descripcion = False
                continue

            # C. CAPTURAR
            if capturando_descripcion:
                # Filtros de limpieza
                if len(txt) < 3: continue
                
                # Chequeo extra de seguridad por si se coló un título
                es_titulo = False
                for bad in blacklist_clean:
                    if bad in txt:
                        es_titulo = True
                        break
                
                if not es_titulo:
                    # Si ya teníamos texto, agregamos salto de línea
                    if ficha_actual["texto_central"]:
                         ficha_actual["texto_central"] += "\n" + txt
                    else:
                         ficha_actual["texto_central"] = txt

        # 4. EXTRAER FOTOS
        sin_fotos = "No se registraron fotografías" in texto_plano_pagina or \
                    "No se registraron fotografias" in texto_plano_pagina or \
                    "No se registraron fotografias" in texto_plano_pagina.lower()

        if not sin_fotos:
            y_titulo_VIII = 0
            tiene_titulo_VIII = False
            
            for b in bloques:
                if "VIII. REGISTRO FOTOGRÁFICO" in b[4]:
                    y_titulo_VIII = b[3]
                    tiene_titulo_VIII = True
                    break

            if len(pagina.get_images()) > 0:
                lista_imagenes = pagina.get_images(full=True)
                for img in lista_imagenes:
                    bbox = pagina.get_image_bbox(img)
                    
                    # FILTRO 1: HEADER/LOGO (<150px)
                    if bbox.y0 < 150: continue
                    
                    # FILTRO 2: Debajo de Título VIII
                    if tiene_titulo_VIII and bbox.y0 < y_titulo_VIII: continue

                    # FILTRO 3: Tamaño
                    base_image = doc.extract_image(img[0])
                    if base_image["width"] < 150 or base_image["height"] < 150: continue
                    
                    # PROCESAR
                    image_bytes = base_image["image"]
                    leyenda_encontrada = ""
                    for b in bloques:
                        b_rect = fitz.Rect(b[:4])
                        b_text = b[4].strip()
                        dist = min(abs(b_rect.y0 - bbox.y1), abs(bbox.y0 - b_rect.y1))
                        if dist < 70 and len(b_text) > 5 and "REGISTRO FOTOGRÁFICO" not in b_text:
                            leyenda_encontrada = b_text
                            break
                    
                    ficha_actual["fotos"].append({
                        "blob": image_bytes,
                        "leyenda": leyenda_encontrada
                    })

    # Guardar última ficha
    if ficha_actual["fecha"] or ficha_actual["texto_central"] or ficha_actual["fotos"]:
        fichas.append(ficha_actual)

    return fichas

# ==========================================
# 3. LÓGICA: GENERADOR EXCEL (DESDE WORD)
# ==========================================

def procesar_word_a_excel(archivo_bytes, nombre_archivo):
    try:
        doc = Document(io.BytesIO(archivo_bytes))
    except Exception as e:
        st.error(f"Error leyendo {nombre_archivo}: {e}")
        return []

    registros = []
    
    for tabla in doc.tables:
        dato = {
            "Fecha": "",
            "Descripción de la actividad": "",
            "Descripción estratigráfica": ""
        }
        encontrado = False 
        
        for fila in tabla.rows:
            for i, celda in enumerate(fila.cells):
                texto_celda = celda.text.strip()
                
                if "Fecha" in texto_celda and len(texto_celda) < 20:
                    if i + 1 < len(fila.cells):
                        dato["Fecha"] = fila.cells[i+1].text.strip()
                        encontrado = True
                
                if "Descripción de la actividad" in texto_celda:
                    if i + 1 < len(fila.cells):
                        dato["Descripción de la actividad"] = fila.cells[i+1].text.strip()
                        encontrado = True

                if "Descripción estratigráfica" in texto_celda:
                    if i + 1 < len(fila.cells):
                        dato["Descripción estratigráfica"] = fila.cells[i+1].text.strip()
                        encontrado = True

        if encontrado:
            if dato["Fecha"] or dato["Descripción de la actividad"] or dato["Descripción estratigráfica"]:
                registros.append(dato)
                
    return registros

# ==========================================
# 4. LÓGICA: GENERADOR FICHAS MAESTRO (DESDE WORD)
# ==========================================

def procesar_maestro_desde_word(archivo_bytes, nombre_archivo):
    try:
        doc = Document(io.BytesIO(archivo_bytes))
    except Exception as e:
        st.error(f"Error leyendo {nombre_archivo}: {e}")
        return []

    fichas = []
    
    for tabla in doc.tables:
        info = {
            "ID Sitio": "", "Coord. Norte": "", "Coord. Este": "", 
            "Categoría": "", "Descripción": "", "Fecha": "", 
            "Responsable": "", "Cronología": "", "foto_blob": None
        }
        es_ficha = False
        crono_checks = [] 
        crono_extra = [] 

        for r_idx, fila in enumerate(tabla.rows):
            for c_idx, celda in enumerate(fila.cells):
                txt = celda.text.strip()
                
                if "ID Sitio" in txt and c_idx + 1 < len(fila.cells):
                    val = fila.cells[c_idx+1].text.strip()
                    if val:
                        info["ID Sitio"] = val
                        es_ficha = True
                
                if "Fecha" in txt and c_idx + 1 < len(fila.cells):
                    info["Fecha"] = fila.cells[c_idx+1].text.strip()
                        
                if "Responsable" in txt and c_idx + 1 < len(fila.cells):
                    info["Responsable"] = fila.cells[c_idx+1].text.strip()

                if "Categoría" in txt and c_idx + 1 < len(fila.cells):
                    info["Categoría"] = fila.cells[c_idx+1].text.strip()

                if "Coord. Central Norte" in txt and c_idx + 1 < len(fila.cells):
                    info["Coord. Norte"] = fila.cells[c_idx+1].text.strip()
                if "Coord. Central Este" in txt and c_idx + 1 < len(fila.cells):
                    info["Coord. Este"] = fila.cells[c_idx+1].text.strip()

                # Descripción
                if txt == "Descripción": 
                    if c_idx + 1 < len(fila.cells):
                        vecino = fila.cells[c_idx+1].text.strip()
                        if "CRONOLOGÍA" not in vecino:
                            info["Descripción"] = vecino
                
                # Cronología
                opciones = ["Prehispánico", "Subactual", "Incierto", "Histórico"]
                for op in opciones:
                    if op in txt:
                        if c_idx + 1 < len(fila.cells):
                            val_vecino = fila.cells[c_idx+1].text.strip().upper()
                            if "X" in val_vecino:
                                crono_checks.append(op)
                
                if "Periodo específico" in txt:
                    if c_idx + 1 < len(fila.cells):
                        val = fila.cells[c_idx+1].text.strip()
                        val = val.replace("Periodo específico:", "").replace("Periodo específico", "").strip()
                        if val and len(val) > 1 and "X" not in val.upper():
                            crono_extra.append(f"Periodo específico: {val}")

                # Foto
                if "Fotografía detalle" in txt:
                    if r_idx > 0:
                        celda_arriba = tabla.rows[r_idx - 1].cells[c_idx]
                        imgs = obtener_imagenes_con_id(celda_arriba._element, doc)
                        if imgs:
                            info["foto_blob"] = imgs[0][1]

        full_crono = crono_checks + crono_extra
        if full_crono:
            info["Cronología"] = ", ".join(list(set(full_crono)))

        if es_ficha and info["ID Sitio"]:
            fichas.append(info)

    return fichas

def crear_doc_tabla_horizontal(datos):
    doc = Document()
    
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    doc.add_heading("Fichas de Hallazgos (Resumen)", 0)
    tabla = doc.add_table(rows=1, cols=9)
    tabla.style = 'Table Grid'
    
    titulos = ["ID Sitio", "Coord. Norte", "Coord. Este", "Cat. (SA/HA)", "Descripción", "Fecha", "Responsable", "Cronología", "Foto"]
    headers = tabla.rows[0].cells
    for i, t in enumerate(titulos):
        headers[i].text = t
        headers[i].paragraphs[0].runs[0].bold = True

    for item in datos:
        row = tabla.add_row().cells
        row[0].text = str(item.get("ID Sitio", ""))
        row[1].text = str(item.get("Coord. Norte", ""))
        row[2].text = str(item.get("Coord. Este", ""))
        row[3].text = str(item.get("Categoría", ""))
        row[4].text = str(item.get("Descripción", ""))
        row[5].text = str(item.get("Fecha", ""))
        row[6].text = str(item.get("Responsable", ""))
        row[7].text = str(item.get("Cronología", ""))
        
        if item.get("foto_blob"):
            p = row[8].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p.add_run()
                run.add_picture(io.BytesIO(item["foto_blob"]), width=Cm(4.5)) 
            except:
                p.add_run("[Err]")
        else:
            row[8].text = "[Sin Foto]"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 5. LÓGICA: GENERADOR KMZ & MAPA INTERACTIVO
# ==========================================

def crear_kml_texto(puntos):
    kml_header = """<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>Hallazgos Arqueológicos</name>"""
    kml_footer = """
  </Document>
</kml>"""
    kml_body = ""
    for p in puntos:
        kml_body += f"""
    <Placemark>
      <name>{p['nombre']}</name>
      <description>{p['desc']}</description>
      <Point>
        <coordinates>{p['lon']},{p['lat']},0</coordinates>
      </Point>
    </Placemark>"""
    return kml_header + kml_body + kml_footer

def obtener_puntos_geograficos_con_foto(archivos):
    """
    Extrae coords y FOTOS para el mapa interactivo.
    """
    try:
        transformer = Transformer.from_crs("epsg:32718", "epsg:4326", always_xy=True)
    except:
        return None

    puntos_acumulados = []
    
    for a in archivos:
        try:
            doc = Document(io.BytesIO(a.read()))
            for tabla in doc.tables:
                id_sitio, norte, este, desc = "", "", "", ""
                foto_bytes = None
                
                for r_idx, fila in enumerate(tabla.rows):
                    for idx, celda in enumerate(fila.cells):
                        txt = celda.text.strip()
                        
                        # Datos
                        if "ID Sitio" in txt and idx+1 < len(fila.cells): id_sitio = fila.cells[idx+1].text.strip()
                        if "Coord. Central Norte" in txt and idx+1 < len(fila.cells): norte = fila.cells[idx+1].text.strip()
                        if "Coord. Central Este" in txt and idx+1 < len(fila.cells): este = fila.cells[idx+1].text.strip()
                        if "Categoría" in txt and idx+1 < len(fila.cells): desc = fila.cells[idx+1].text.strip()
                        
                        # Foto (para el mapa)
                        if "Fotografía detalle" in txt and r_idx > 0:
                            celda_arriba = tabla.rows[r_idx - 1].cells[idx]
                            imgs = obtener_imagenes_con_id(celda_arriba._element, doc)
                            if imgs:
                                foto_bytes = imgs[0][1]
                
                if id_sitio and norte and este:
                    n = limpiar_coordenada(norte)
                    e = limpiar_coordenada(este)
                    if n and e:
                        lon, lat = transformer.transform(e, n)
                        puntos_acumulados.append({
                            "nombre": id_sitio, 
                            "desc": desc, 
                            "lat": lat, 
                            "lon": lon,
                            "foto": foto_bytes
                        })
        except:
            continue
            
    return puntos_acumulados

# ==========================================
#          MENÚ LATERAL
# ==========================================

st.sidebar.title("Arqueología App")
opcion = st.sidebar.radio("Herramientas:", [
    "Generador Word (MAP)", 
    "Generador Word MAP (Desde PDF)", 
    "Generador Excel (Desde Word)",
    "Generador Fichas (Desde Word)",
    "Generador KMZ (Georreferenciación)",
    "Visor de Mapa Interactivo"
])

# 1. Generador Word (MAP - Desde Word)
if opcion == "Generador Word (MAP)":
    st.title("Generador Word MAP (Desde DOCX)")
    st.markdown("Crea la tabla resumen mensual a partir de los anexos diarios en Word.")
    st.info("Configuración: Franklin Gothic Book 9 | Fotos 8x6 cm | Centrado")
    archivos = st.file_uploader("Subir Anexos Word (.docx)", accept_multiple_files=True, key="word_up")
    if archivos and st.button("Generar Informe Word"):
        todas = []
        bar = st.progress(0)
        for i, a in enumerate(archivos):
            fichas = procesar_archivo_v12(a.read(), a.name)
            todas.extend(fichas)
            bar.progress((i+1)/len(archivos))
        if todas:
            todas.sort(key=lambda x: x['fecha'] if x['fecha'] else "ZZZ")
            doc_out = generar_word_con_formato(todas)
            st.success("✅ Informe Word generado.")
            st.download_button("Descargar Word", doc_out, "Resumen_MAP.docx")
        else: st.error("No se encontraron datos.")

# 1.1 Generador Word MAP (Desde PDF)
elif opcion == "Generador Word MAP (Desde PDF)":
    st.title("Generador Word MAP (Desde PDF)")
    st.markdown("Crea la tabla resumen mensual extrayendo datos de reportes en PDF.")
    st.warning("Requiere librería 'pymupdf' instalada.")
    
    archivos = st.file_uploader("Subir Reportes PDF (.pdf)", accept_multiple_files=True, key="pdf_up")
    
    if archivos and st.button("Procesar PDFs y Generar Word"):
        todas_fichas = []
        bar = st.progress(0)
        
        for i, a in enumerate(archivos):
            fichas = procesar_pdf_a_word_map(a.read(), a.name)
            todas_fichas.extend(fichas)
            bar.progress((i+1)/len(archivos))
            
        if todas_fichas:
            # Ordenar por fecha si es posible
            todas_fichas.sort(key=lambda x: x['fecha'] if x['fecha'] else "ZZZ")
            
            # Reutilizamos la función de formato que ya existe
            doc_out = generar_word_con_formato(todas_fichas)
            
            st.success(f"✅ Se procesaron {len(todas_fichas)} fichas desde PDF.")
            st.download_button("Descargar Word Resumen", doc_out, "Resumen_MAP_Desde_PDF.docx")
        else:
            st.error("No se pudieron extraer datos válidos de los PDFs.")

# 2. Generador Excel (Desde Word)
elif opcion == "Generador Excel (Desde Word)":
    st.title("Generador Excel (Desde Word)")
    st.markdown("Extrae: Fecha, Descripción de actividad y estratigráfica (celda vecina).")
    archivos = st.file_uploader("Subir Anexos Word (.docx)", accept_multiple_files=True, key="word_excel_up")
    if archivos and st.button("Generar Excel"):
        todos_registros = []
        bar = st.progress(0)
        for i, a in enumerate(archivos):
            regs = procesar_word_a_excel(a.read(), a.name)
            todos_registros.extend(regs)
            bar.progress((i+1)/len(archivos))
        if todos_registros:
            df = pd.DataFrame(todos_registros)
            st.success(f"✅ Se extrajeron {len(df)} filas.")
            st.dataframe(df)
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name="Resumen")
            st.download_button("⬇️ Descargar Excel", buffer.getvalue(), "Resumen_Word_Excel.xlsx")
        else: st.error("No se encontraron datos.")

# 3. Generador Fichas (Desde Word)
elif opcion == "Generador Fichas (Desde Word)":
    st.title("Generador de Fichas (Desde DOCX)")
    st.markdown("Extrae datos y fotos desde las Fichas de Hallazgo originales en Word.")
    archivos = st.file_uploader("Subir Fichas de Hallazgo (.docx)", accept_multiple_files=True, key="maestro_up")
    if archivos and st.button("Procesar Archivos"):
        todos_datos = []
        bar = st.progress(0)
        for i, a in enumerate(archivos):
            datos = procesar_maestro_desde_word(a.read(), a.name)
            todos_datos.extend(datos)
            bar.progress((i+1)/len(archivos))
        if todos_datos:
            st.success(f"✅ Se procesaron {len(todos_datos)} fichas.")
            df = pd.DataFrame(todos_datos)
            df_excel = df.drop(columns=["foto_blob"], errors='ignore')
            orden = ["ID Sitio", "Coord. Norte", "Coord. Este", "Categoría", "Descripción", "Fecha", "Responsable", "Cronología"]
            cols = [c for c in orden if c in df_excel.columns]
            df_excel = df_excel[cols]
            buf_excel = io.BytesIO()
            with pd.ExcelWriter(buf_excel, engine='openpyxl') as writer:
                df_excel.to_excel(writer, index=False, sheet_name="Hallazgos")
            buf_word = crear_doc_tabla_horizontal(todos_datos)
            col1, col2 = st.columns(2)
            col1.download_button("⬇️ Descargar Excel", buf_excel.getvalue(), "Base_Datos_Hallazgos.xlsx")
            col2.download_button("⬇️ Descargar Fichas Word", buf_word.getvalue(), "Fichas_Con_Fotos.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.dataframe(df_excel)
        else: st.error("No se encontraron fichas válidas.")

# 4. Generador KMZ
elif opcion == "Generador KMZ (Georreferenciación)":
    st.title("Generador KMZ (Google Earth)")
    st.markdown("Crea un archivo KMZ a partir de las coordenadas (UTM 18S) en los documentos Word.")
    archivos = st.file_uploader("Subir Fichas de Hallazgo (.docx)", accept_multiple_files=True, key="kmz_up")
    if archivos and st.button("Generar KMZ"):
        try:
            puntos = obtener_puntos_geograficos_con_foto(archivos)
            if puntos:
                kmz_final_buffer = io.BytesIO()
                with zipfile.ZipFile(kmz_final_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    kml_content = crear_kml_texto(puntos)
                    zf.writestr("doc.kml", kml_content)
                kmz_final_buffer.seek(0)
                st.success(f"✅ Se generaron {len(puntos)} puntos.")
                st.download_button("⬇️ Descargar KMZ", kmz_final_buffer.getvalue(), "Hallazgos_Georreferenciados.kmz")
            else: st.error("No se encontraron coordenadas válidas.")
        except ImportError: st.error("Falta librería 'pyproj'.")

# 5. Visor Mapa Interactivo
elif opcion == "Visor de Mapa Interactivo":
    st.title("Visor de Mapa Interactivo")
    st.markdown("Visualiza los hallazgos en Google Satélite con fotos.")
    
    try:
        import folium
        from streamlit_folium import st_folium
    except ImportError:
        st.error("⚠️ Faltan las librerías 'folium' y 'streamlit-folium'.")
        st.stop()

    archivos = st.file_uploader("Subir Fichas de Hallazgo (.docx)", accept_multiple_files=True, key="mapa_up")
    
    if 'map_points' not in st.session_state:
        st.session_state.map_points = None

    if archivos and st.button("Procesar y Mostrar Mapa"):
        with st.spinner("Leyendo coordenadas y fotos..."):
            puntos = obtener_puntos_geograficos_con_foto(archivos)
            if puntos:
                st.session_state.map_points = puntos
            else:
                st.error("No se pudieron extraer datos.")

    if st.session_state.map_points:
        puntos = st.session_state.map_points
        st.success(f"✅ Se encontraron {len(puntos)} puntos.")
        
        avg_lat = sum(p['lat'] for p in puntos) / len(puntos)
        avg_lon = sum(p['lon'] for p in puntos) / len(puntos)
        
        # Mapa base limpio para poner Google Sat
        m = folium.Map(location=[avg_lat, avg_lon], zoom_start=12, tiles=None)
        
        # Capa Satélite
        folium.TileLayer(
            tiles='https://mt1.google.com/vt/lyrs=s&x={x}&y={y}&z={z}',
            attr='Google',
            name='Google Satellite',
            overlay=False,
            control=True
        ).add_to(m)
        
        for p in puntos:
            # HTML Popup con Foto Base64
            html = f"<div style='font-family: Arial; width: 200px;'>"
            html += f"<b>{p['nombre']}</b><br><i style='font-size:12px'>{p['desc']}</i>"
            
            if p['foto']:
                # Convertir bytes a base64
                b64 = base64.b64encode(p['foto']).decode('utf-8')
                html += f"<br><img src='data:image/jpeg;base64,{b64}' width='100%' style='margin-top:5px; border-radius:5px;'>"
            
            html += "</div>"
            
            iframe = folium.IFrame(html, width=220, height=220)
            popup = folium.Popup(iframe, max_width=220)
            
            # Marcador como PUNTO ROJO
            folium.CircleMarker(
                location=[p['lat'], p['lon']],
                radius=6,
                color='red',
                fill=True,
                fill_color='red',
                fill_opacity=1.0,
                popup=popup,
                tooltip=p['nombre']
            ).add_to(m)
        
        st_folium(m, width=900, height=600)
